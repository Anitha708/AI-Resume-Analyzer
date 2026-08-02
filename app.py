import streamlit as st
import pdfplumber
from docx import Document
import re

# --------------------------------------------------
# PAGE CONFIG
# --------------------------------------------------

st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="📄",
    layout="wide"
)

st.title("📄 AI Resume Analyzer")
st.write("Upload any PDF or DOCX resume to extract and analyze its content.")

# --------------------------------------------------
# SKILLS DATABASE
# --------------------------------------------------

SKILLS = {
    "Programming": [
        "python",
        "java",
        "c++",
        "c",
        "javascript",
        "typescript",
        "r",
        "matlab"
    ],

    "AI & Machine Learning": [
        "artificial intelligence",
        "machine learning",
        "deep learning",
        "machine learning algorithms",
        "generative ai",
        "natural language processing",
        "nlp",
        "computer vision",
        "tensorflow",
        "keras",
        "pytorch",
        "scikit-learn"
    ],

    "Data & Analytics": [
        "data analysis",
        "data analytics",
        "exploratory data analysis",
        "eda",
        "pandas",
        "numpy",
        "matplotlib",
        "seaborn",
        "power bi",
        "tableau",
        "sql",
        "excel"
    ],

    "Web & Development": [
        "html",
        "css",
        "react",
        "node.js",
        "flask",
        "django",
        "streamlit"
    ],

    "Cloud & DevOps": [
        "aws",
        "azure",
        "google cloud",
        "gcp",
        "docker",
        "kubernetes",
        "git",
        "github",
        "gitlab"
    ]
}

# --------------------------------------------------
# EXTRACT PDF TEXT
# --------------------------------------------------

def extract_pdf_text(uploaded_file):

    text = ""

    try:
        with pdfplumber.open(uploaded_file) as pdf:

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"

    except Exception as e:
        st.error(f"❌ Error reading PDF: {e}")
        return ""

    return text


# --------------------------------------------------
# EXTRACT DOCX TEXT
# --------------------------------------------------

def extract_docx_text(uploaded_file):

    text = ""

    try:
        document = Document(uploaded_file)

        # Paragraphs
        for paragraph in document.paragraphs:

            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Tables
        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():
                        text += cell.text + "\n"

    except Exception as e:
        st.error(f"❌ Error reading DOCX: {e}")
        return ""

    return text


# --------------------------------------------------
# CLEAN TEXT
# --------------------------------------------------

def clean_text(text):

    text = text.lower()

    # Replace multiple spaces
    text = re.sub(r"\s+", " ", text)

    # Remove unnecessary special characters
    text = re.sub(r"[^\w\s+#.-]", " ", text)

    # Remove extra spaces
    text = re.sub(r"\s+", " ", text)

    return text.strip()


# --------------------------------------------------
# FIND SKILLS
# --------------------------------------------------

def find_skills(text):

    cleaned_text = clean_text(text)

    found_skills = {}

    for category, skills in SKILLS.items():

        category_skills = []

        for skill in skills:

            # Special handling for skills containing symbols
            pattern = r"(?<!\w)" + re.escape(skill.lower()) + r"(?!\w)"

            if re.search(pattern, cleaned_text):
                category_skills.append(skill.title())

        if category_skills:
            found_skills[category] = category_skills

    return found_skills


# --------------------------------------------------
# FIND RESUME SECTIONS
# --------------------------------------------------

def find_sections(text):

    cleaned = clean_text(text)

    sections = {
        "Professional Summary": [
            "professional summary",
            "summary",
            "profile",
            "objective"
        ],

        "Education": [
            "education",
            "academic background"
        ],

        "Experience": [
            "experience",
            "work experience",
            "employment"
        ],

        "Internships": [
            "internships",
            "internship"
        ],

        "Projects": [
            "projects",
            "project"
        ],

        "Technical Skills": [
            "technical skills",
            "skills",
            "technical expertise"
        ],

        "Certifications": [
            "certifications",
            "certificates",
            "certification"
        ],

        "Languages": [
            "languages",
            "language proficiency"
        ]
    }

    detected = []

    for section, keywords in sections.items():

        for keyword in keywords:

            if keyword in cleaned:
                detected.append(section)
                break

    return detected


# --------------------------------------------------
# RESUME SCORE
# --------------------------------------------------

def calculate_resume_score(text, sections, skills_found):
    score = 0
    text_lower = text.lower()

    # 1. Professional Summary - 10 points
    if "professional summary" in text_lower or "summary" in text_lower:
        score += 10

    # 2. Education - 10 points
    if "education" in text_lower:
        score += 10

    # 3. Experience / Internships - 15 points
    if "experience" in text_lower or "internships" in text_lower:
        score += 15

    # 4. Projects - 15 points
    if "projects" in text_lower:
        score += 15

    # 5. Technical Skills - 15 points
    if "technical skills" in text_lower or "skills" in text_lower:
        score += 15

    # 6. Certifications - 10 points
    if "certifications" in text_lower or "certification" in text_lower:
        score += 10

    # 7. Contact information - 10 points
    if "@" in text and any(char.isdigit() for char in text):
        score += 10

    # 8. Relevant skills - 10 points
    if len(skills_found) >= 10:
        score += 10
    elif len(skills_found) >= 5:
        score += 7
    elif len(skills_found) >= 2:
        score += 4

    # 9. Resume length / content - 5 points
    word_count = len(text.split())

    if word_count >= 300:
        score += 5
    elif word_count >= 150:
        score += 3

    return min(score, 100)

# --------------------------------------------------
# FILE UPLOADER
# --------------------------------------------------

uploaded_file = st.file_uploader(
    "Upload your Resume",
    type=["pdf", "docx"],
    key="resume_upload"
)


# --------------------------------------------------
# PROCESS UPLOADED FILE
# --------------------------------------------------

if uploaded_file is not None:

    st.success(f"Uploaded: {uploaded_file.name}")

    # Important:
    # Reset file pointer before reading
    uploaded_file.seek(0)

    # PDF
    if uploaded_file.name.lower().endswith(".pdf"):

        extracted_text = extract_pdf_text(uploaded_file)

    # DOCX
    elif uploaded_file.name.lower().endswith(".docx"):

        extracted_text = extract_docx_text(uploaded_file)

    else:

        extracted_text = ""

    # --------------------------------------------------
    # CHECK TEXT
    # --------------------------------------------------

    if not extracted_text.strip():

        st.error(
            "❌ No readable text was found in this resume. "
            "Please upload a text-based PDF or DOCX file."
        )

        st.stop()

    # --------------------------------------------------
    # EXTRACTED TEXT
    # --------------------------------------------------

    st.subheader("📋 Extracted Resume Text")

    with st.expander("View Extracted Text", expanded=True):

        st.text_area(
            "Resume Content",
            extracted_text,
            height=400
        )

    # --------------------------------------------------
    # CLEAN TEXT
    # --------------------------------------------------

    cleaned_text = clean_text(extracted_text)

    st.subheader("🧹 Cleaned Resume Text")

    with st.expander("View Cleaned Text"):

        st.text_area(
            "Processed Content",
            cleaned_text,
            height=300
        )

    # --------------------------------------------------
    # SKILLS
    # --------------------------------------------------

    found_skills = find_skills(extracted_text)

    total_skills = sum(
        len(skills)
        for skills in found_skills.values()
    )

    # --------------------------------------------------
    # SECTIONS
    # --------------------------------------------------

    detected_sections = find_sections(extracted_text)

    # --------------------------------------------------
    # SCORE
    # --------------------------------------------------

    score = calculate_resume_score(
    extracted_text,
    detected_sections,
    found_skills
)

    # --------------------------------------------------
    # SCORE DISPLAY
    # --------------------------------------------------

    st.subheader("📊 Resume Score")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("Resume Score", f"{score}/100")

    with col2:
        st.metric("Skills Found", total_skills)

    with col3:
        st.metric(
            "Sections Found",
            len(detected_sections)
        )

    # --------------------------------------------------
    # SKILLS DISPLAY
    # --------------------------------------------------

    st.subheader("🛠️ Skills Identified")

    if found_skills:

        for category, skills in found_skills.items():

            st.markdown(f"### {category}")

            for skill in skills:

                st.write(f"• {skill}")

    else:

        st.warning(
            "No matching technical skills were detected."
        )

    # --------------------------------------------------
    # SECTIONS
    # --------------------------------------------------

    st.subheader("📑 Resume Sections Detected")

    all_sections = [
        "Professional Summary",
        "Education",
        "Experience",
        "Internships",
        "Projects",
        "Technical Skills",
        "Certifications",
        "Languages"
    ]

    for section in all_sections:

        if section in detected_sections:

            st.success(f"✓ {section}")

        else:

            st.write(f"○ {section} — Not detected")

    # --------------------------------------------------
    # RECOMMENDATIONS
    # --------------------------------------------------

    

    st.subheader("💡 Resume Recommendations")

    recommendations = []

    text_lower = extracted_text.lower()

    if "summary" not in text_lower:
        recommendations.append(
            "Add a professional summary describing your career goal and key strengths."
        )

    if "education" not in text_lower:
        recommendations.append(
            "Add an Education section with your degree, college, and graduation year."
        )

    if "experience" not in text_lower and "internship" not in text_lower:
        recommendations.append(
            "Add an Experience or Internship section with your role, organization, dates, and responsibilities."
        )

    if "projects" not in text_lower:
        recommendations.append(
            "Add 2–3 relevant projects and mention the technologies used."
        )

    if "certification" not in text_lower:
        recommendations.append(
            "Add relevant certifications to strengthen your profile."
        )

    if "skill" not in text_lower:
        recommendations.append(
            "Add a Technical Skills section."
        )

    if total_skills < 5:
        recommendations.append(
            "Add more relevant technical skills that you have actually used."
        )

    elif total_skills < 10:
        recommendations.append(
            "Consider adding more relevant technical skills to strengthen your profile."
        )

    if "github" not in text_lower:
        recommendations.append(
            "Add your GitHub profile and relevant project repositories."
        )

    if recommendations:
        for recommendation in recommendations:
            st.warning("⚠️ " + recommendation)
    else:
        st.success(
            "Excellent! Your resume contains the major sections and relevant skills expected for an entry-level technical resume."
        )